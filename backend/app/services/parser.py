"""Resume file parser.

Supports PDF (via PyMuPDF / fitz) and DOCX (via python-docx).
Entry point: ``parse_resume(path)`` dispatches based on file extension.
"""

from __future__ import annotations

import json
import logging
import os
import re
from dataclasses import dataclass, field
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)

_EMAIL_RE = re.compile(r'\b[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}\b')
_PHONE_RE = re.compile(
    r'(?:\+?\d{1,3}[\s\-.]?)?'         # optional country code
    r'(?:\(?\d{2,4}\)?[\s\-.]?)?'       # optional area code
    r'\d{3,5}[\s\-.]?\d{3,5}'           # main number
)

_EXEC_RE = re.compile(
    r'\b(c[tfl]o|ceo|coo|cpo|vice\s+president|vp\b|head\s+of|chief\b)',
    re.I,
)
_SENIOR_RE = re.compile(
    r'\b(senior|sr\.?\s|lead|principal|staff\s+\w+|architect|director|manager)\b',
    re.I,
)
_ENTRY_RE = re.compile(
    r'\b(intern\b|internship|trainee|fresher|fresh\s+graduate|recent\s+graduate|'
    r'graduate\s+engineer|engineering\s+student|b\.?tech\s+student|final\s+year|passout|pass\s*out)\b',
    re.I,
)
_JUNIOR_RE = re.compile(
    r'\b(junior|jr\.?\s|entry[\s\-]level)\b',
    re.I,
)
_TITLE_WORDS = {
    'engineer', 'developer', 'analyst', 'manager', 'designer', 'architect',
    'lead', 'director', 'consultant', 'specialist', 'scientist', 'officer',
    'executive', 'intern', 'associate', 'programmer', 'administrator',
    'technician', 'strategist', 'coordinator', 'head',
}
_DATE_LINE_RE = re.compile(r'^[\d\s\-/,–—|•]+$')

# Matches internship/trainee role indicators — used to exclude these from years_experience
_INTERN_KEYWORDS_RE = re.compile(
    r'\b(intern(?:ship)?s?|trainee|apprentice|student\s+(?:engineer|developer|researcher|worker)|'
    r'industrial\s+train(?:ing|ee)|summer\s+(?:project|training|program))\b',
    re.I,
)

# Matches: "2018 - 2022", "Jan 2020 – Mar 2023", "10/2020 – 05/2024", "2021 - Present"
# Optional MM/ or DD/ prefix before the 4-digit year handles MM/YYYY format.
_DATE_RANGE_RE = re.compile(
    r'(?:(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+|\d{1,2}[/\-])?'
    r'(\d{4})'
    r'\s*[-–—to]+\s*'
    r'((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4}'
    r'|(?:\d{1,2}[/\-])?\d{4}'
    r'|present|current|now|today)',
    re.I,
)


def _parse_year(token: str) -> int:
    """Extract a 4-digit year from a token like '2023', 'Mar 2023', 'present'."""
    if re.match(r'^(present|current|now|today)$', token.strip(), re.I):
        return datetime.now().year
    m = re.search(r'\d{4}', token)
    return int(m.group()) if m else datetime.now().year


def _is_internship_context(text: str, match_start: int) -> bool:
    """Return True if the 5 lines preceding a date-range match contain internship keywords."""
    preceding = text[max(0, match_start - 400):match_start]
    lines = [l.strip() for l in preceding.splitlines() if l.strip()]
    context = "\n".join(lines[-5:]) if lines else ""
    return bool(_INTERN_KEYWORDS_RE.search(context))


def _merge_intervals(intervals: List[Tuple[int, int]]) -> float:
    """Sum non-overlapping year spans from a list of (start, end) tuples."""
    if not intervals:
        return 0.0
    intervals.sort()
    merged: List[Tuple[int, int]] = [intervals[0]]
    for start, end in intervals[1:]:
        if start <= merged[-1][1]:
            merged[-1] = (merged[-1][0], max(merged[-1][1], end))
        else:
            merged.append((start, end))
    return sum(max(0.0, end - start) for start, end in merged)


# Matches "5+ years", "5 years", "over 5 years of experience", "5-6 years"
_YEARS_TEXT_RE = re.compile(
    r'(?:over\s+|more\s+than\s+|approximately\s+)?(\d+(?:\.\d+)?)\s*\+?\s*(?:to|-)\s*(?:\d+\s+)?years?'
    r'|(?:over\s+|more\s+than\s+)?(\d+(?:\.\d+)?)\s*\+\s*years?',
    re.I,
)


def extract_years_experience(sections: "List[Any]") -> Optional[float]:
    """Calculate total years of experience.

    Primary: sum non-overlapping date ranges from experience sections.
    Fallback: parse explicit "N years" / "N+ years" phrases from summary sections.
    """
    intervals: List[Tuple[int, int]] = []
    summary_text = ""
    for sec in sections:
        sec_type = sec.type if hasattr(sec, "type") else sec.get("type", "")
        text = sec.text if hasattr(sec, "text") else sec.get("text", "")
        if sec_type in ("experience", "work_experience"):
            for m in _DATE_RANGE_RE.finditer(text):
                if _is_internship_context(text, m.start()):
                    continue
                start = _parse_year(m.group(1))
                end = _parse_year(m.group(2))
                if 1970 <= start <= datetime.now().year and start <= end <= datetime.now().year + 1:
                    intervals.append((start, end))
        elif sec_type in ("summary", "profile", "objective", "about"):
            summary_text += " " + text

    if intervals:
        return round(_merge_intervals(intervals), 1)

    # Fallback: explicit mention in summary/profile text
    for m in _YEARS_TEXT_RE.finditer(summary_text):
        val = float(m.group(1) or m.group(2))
        if 0 < val < 50:
            return val

    return None


# Matches standalone 4-digit years in education context (e.g. "B.Tech 2023", "2024")
_STANDALONE_YEAR_RE = re.compile(r'\b(20\d{2}|19[89]\d)\b')
# Lines containing degree keywords — used as raw-text fallback for graduation year
_DEGREE_LINE_RE = re.compile(
    r'\b(bachelor|b\.?\s*tech|b\.?\s*e\.?|b\.?\s*sc|master|m\.?\s*tech|m\.?\s*e\.?|m\.?\s*sc'
    r'|ph\.?\s*d|mba|diploma|b\.?\s*com|bca|mca|b\.?\s*a\.?|associate\s+degree'
    r'|b\.?\s*c\.?a\.?|m\.?\s*c\.?a\.?|b\.?\s*c\.?s\.?|engineering|polytechnic|graduation)\b',
    re.I,
)
# Matches "Expected 2025", "2025 (expected)", "Graduating 2025", "Class of 2024",
# "Batch of 2023", "Passed out: 2023", "Passout 2024", "Year of Passing: 2024"
_GRAD_YEAR_RE = re.compile(
    r'(?:expected|graduating|graduation|class\s+of|batch(?:\s+of)?|year\s+of\s+pass(?:ing)?'
    r'|pass(?:ed)?\s*(?:out)?\s*[:\-]?|passout)\s*(20\d{2}|19[89]\d)'
    r'|(20\d{2}|19[89]\d)\s*\(expected\)',
    re.I,
)
# Matches abbreviated year ranges like "2020-22" or "Aug 2020 - Jun 22"
_ABBREV_YEAR_RE = re.compile(
    r'(?:(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+)?'
    r'(20\d{2}|19[89]\d)'       # start year (4 digits)
    r'\s*[-–—to]+\s*'
    r'(?:(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+)?'
    r'(?!(present|current|now|today))'  # not "present"
    r'(\d{2})\b',                        # abbreviated end year e.g. "22"
    re.I,
)
# Matches parenthesized date ranges like "(2019-2023)" or "(Aug 2019 – May 2023)"
_PAREN_DATE_RE = re.compile(
    r'\(?'
    r'(?:(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+)?'
    r'(20\d{2}|19[89]\d)'
    r'\s*[-–—to]+\s*'
    r'(?:(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+)?'
    r'(20\d{2}|19[89]\d)'
    r'\)?',
    re.I,
)


def extract_graduation_year(sections: "List[Any]", raw_text: str = "") -> Optional[int]:
    """Return the most recent education graduation year.

    Checks education sections first (by multiple possible section type names).
    Falls back to scanning raw_text lines that contain degree keywords, checking
    up to 5 following lines for associated date ranges.

    Handles:
    - Standard date ranges: "2020 – 2024", "Aug 2020 – May 2024"
    - Abbreviated end years: "2020-22" → 2022
    - Parenthesized ranges: "(2019-2023)"
    - Expected/batch/passout keywords
    - CGPA-adjacent years: "CGPA 8.5 | 2024"
    - Accepts years from 1980 up to current year + 2 (upcoming graduates).
    """
    current_year = datetime.now().year
    max_year = current_year + 2
    candidates: list[int] = []

    # All section type names that indicate education
    _EDUCATION_TYPES = {
        "education", "education_background", "educational", "educational_background",
        "education_details", "academic", "academic_background", "academics",
    }

    def _resolve_abbrev_year(start_yr: int, abbrev: str) -> int:
        """Expand a 2-digit abbreviated year relative to a 4-digit start year.
        e.g. start=2020, abbrev='22' → 2022
        """
        century = (start_yr // 100) * 100
        full = century + int(abbrev)
        # If abbreviated year is less than last 2 digits of start, it wraps to next century
        if full < start_yr:
            full += 100
        return full

    def _collect_years(text: str) -> None:
        # 1. Explicit graduation keywords (highest confidence)
        for m in _GRAD_YEAR_RE.finditer(text):
            yr = int(m.group(1) or m.group(2))
            if 1980 <= yr <= max_year:
                candidates.append(yr)
        # 2. Parenthesized date ranges like (2019-2023)
        for m in _PAREN_DATE_RE.finditer(text):
            yr = int(m.group(2))
            if 1980 <= yr <= max_year:
                candidates.append(yr)
        # 3. Standard date ranges — take the END year
        for m in _DATE_RANGE_RE.finditer(text):
            end_token = m.group(2)
            # Skip "present/current" to avoid inflating to current year
            if re.match(r'^(present|current|now|today)$', end_token.strip(), re.I):
                continue
            end_yr = _parse_year(end_token)
            if 1980 <= end_yr <= max_year:
                candidates.append(end_yr)
        # 4. Abbreviated year ranges like "2020-22"
        for m in _ABBREV_YEAR_RE.finditer(text):
            start_yr = int(m.group(1))
            full_end = _resolve_abbrev_year(start_yr, m.group(2))
            if 1980 <= full_end <= max_year:
                candidates.append(full_end)
        # 5. Standalone 4-digit years (lowest priority — catches e.g. "B.Tech 2024")
        for m in _STANDALONE_YEAR_RE.finditer(text):
            yr = int(m.group(1))
            if 1980 <= yr <= max_year:
                candidates.append(yr)

    # ── Primary: scan structured education sections ───────────────────────────
    for sec in sections:
        sec_type = sec.type if hasattr(sec, "type") else sec.get("type", "")
        if sec_type.lower() not in _EDUCATION_TYPES:
            continue
        _collect_years(sec.text if hasattr(sec, "text") else sec.get("text", ""))

    # ── Fallback: scan raw text lines containing degree keywords ─────────────
    # Catches cases where PDF right-aligned dates aren't in the education section.
    if raw_text:
        lines = raw_text.splitlines()
        for i, line in enumerate(lines):
            if not _DEGREE_LINE_RE.search(line):
                continue
            # Collect from the degree line itself
            _collect_years(line)
            # Check up to 5 following lines for associated date ranges
            for adjacent in lines[i + 1: i + 6]:
                adj_lower = adjacent.lower()
                # Stop if we've hit a clearly different section heading
                if any(kw in adj_lower for kw in (
                    "experience", "project", "skill", "objective",
                    "certification", "achievement", "award", "publication",
                )):
                    break
                if (
                    _DATE_RANGE_RE.search(adjacent)
                    or _ABBREV_YEAR_RE.search(adjacent)
                    or _PAREN_DATE_RE.search(adjacent)
                    or _STANDALONE_YEAR_RE.search(adjacent)
                ):
                    _collect_years(adjacent)

    return max(candidates) if candidates else None


def infer_experience_level(text: str) -> str:
    """Keyword-only fallback when years_experience cannot be extracted.

    entry=interns/freshers, junior=jr. titles, senior/executive by title keywords.
    Defaults to 'entry' when nothing matches.
    """
    sample = text[:6000]
    if _ENTRY_RE.search(sample):
        return "entry"
    if _JUNIOR_RE.search(sample):
        return "junior"
    if _EXEC_RE.search(sample):
        return "executive"
    if _SENIOR_RE.search(sample):
        return "senior"
    return "entry"


def extract_current_title(sections: "List[Any]") -> Optional[str]:
    """Return most recent job title from a list of Section objects or dicts."""
    for sec in sections:
        sec_type = sec.type if hasattr(sec, "type") else sec.get("type", "")
        if sec_type != "experience":
            continue
        raw_text = sec.text if hasattr(sec, "text") else sec.get("text", "")
        for line in raw_text.splitlines():
            line = line.strip()
            if not line or len(line) < 4:
                continue
            if _DATE_LINE_RE.match(line):
                continue
            if len(line) > 120:
                continue
            words = set(line.lower().split())
            if words & _TITLE_WORDS:
                return line[:120]
            if 2 <= len(line.split()) <= 7:
                return line[:120]
    return None


def extract_phone(text: str) -> Optional[str]:
    """Return first phone number found in the first 2000 chars of text."""
    header = text[:2000]
    # Look for phone patterns, skip numbers that are too short (likely years/IDs)
    for m in _PHONE_RE.finditer(header):
        digits = re.sub(r'\D', '', m.group(0))
        if 7 <= len(digits) <= 15:  # valid phone number length
            return m.group(0).strip()
    return None


def extract_email(text: str) -> Optional[str]:
    """Return first email found in text (header area preferred)."""
    header = text[:2000]
    m = _EMAIL_RE.search(header)
    if m:
        return m.group(0).lower()
    m = _EMAIL_RE.search(text)
    return m.group(0).lower() if m else None


_SECTION_HEADER_BLACKLIST = {
    "career objective",
    "professional summary",
    "work experience",
    "education background",
    "educational background",
    "skills summary",
    "technical skills",
    "core competencies",
    "key skills",
    "areas of expertise",
    "summary of qualifications",
    "personal information",
    "contact information",
    "references available",
    "about me",
    "profile summary",
    "career summary",
    "objective statement",
    "professional profile",
}

# Words that cannot appear in a person's name — used by extract_name to reject skill/section lines
_NON_NAME_WORDS = {
    # Document meta / field labels — these appear as "Name: John Smith" prefixes
    'resume', 'cv', 'curriculum', 'vitae', 'page', 'profile',
    'contact', 'email', 'phone', 'mobile', 'tel', 'fax',
    'address', 'website', 'portfolio', 'linkedin', 'github',
    'name', 'candidate', 'applicant', 'subject',
    'location', 'city', 'state', 'country', 'nationality', 'gender',
    'date', 'birth', 'dob', 'age', 'marital', 'status',
    'declaration', 'signature', 'place',
    # Common section headings (individual words)
    'skills', 'technical', 'experience', 'education', 'projects',
    'summary', 'objective', 'achievements', 'certifications', 'interests',
    'languages', 'programming', 'technologies', 'tools', 'frameworks',
    'hobbies', 'references', 'highlights', 'awards', 'publications',
    'qualifications', 'competencies', 'expertise', 'proficiencies',
    # Programming languages
    'python', 'java', 'javascript', 'typescript', 'kotlin', 'swift',
    'ruby', 'php', 'golang', 'rust', 'scala', 'perl', 'matlab',
    'haskell', 'dart', 'elixir', 'lua', 'groovy', 'clojure', 'cobol',
    # Tech ecosystems & platforms
    'linux', 'windows', 'android', 'macos', 'ios', 'unix',
    'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform',
    'react', 'angular', 'vue', 'node', 'django', 'flask',
    'spring', 'bootstrap', 'jquery', 'express', 'fastapi', 'laravel',
    'html', 'css', 'sql', 'mysql', 'postgres', 'mongodb', 'redis',
    'git', 'github', 'gitlab', 'jira', 'confluence', 'jenkins',
    # General tech nouns
    'software', 'hardware', 'system', 'systems', 'database', 'databases',
    'computer', 'science', 'technology', 'engineering', 'development',
    'backend', 'frontend', 'fullstack', 'devops', 'cloud', 'security',
    # Education nouns
    'bachelor', 'master', 'degree', 'university', 'college', 'institute',
    'school', 'department', 'faculty', 'graduation', 'honours',
    # Job titles / role words that should never appear in a person's name
    'intern', 'internship', 'engineer', 'developer', 'designer', 'analyst',
    'manager', 'director', 'executive', 'officer', 'consultant', 'architect',
    'specialist', 'coordinator', 'administrator', 'associate', 'assistant',
    'lead', 'senior', 'junior', 'principal', 'head', 'chief', 'officer',
    'product', 'business', 'data', 'application', 'professional', 'trainee',
    'researcher', 'scientist', 'strategist', 'operations', 'marketing',
    'sales', 'finance', 'accounting', 'legal', 'compliance', 'support',
    # Geographic — Indian cities, states, and countries that appear in address lines
    # and must never be mistaken for a person's name (e.g. "Hyderabad, Telangana")
    'hyderabad', 'telangana', 'bangalore', 'bengaluru', 'mumbai', 'delhi',
    'chennai', 'kolkata', 'calcutta', 'pune', 'ahmedabad', 'jaipur',
    'lucknow', 'kanpur', 'nagpur', 'indore', 'thane', 'bhopal', 'visakhapatnam',
    'patna', 'vadodara', 'ghaziabad', 'ludhiana', 'agra', 'nashik', 'ranchi',
    'surat', 'coimbatore', 'kochi', 'vijayawada', 'rajkot', 'meerut',
    'india', 'karnataka', 'maharashtra', 'gujarat', 'rajasthan',
    'andhra', 'pradesh', 'kerala', 'haryana', 'punjab', 'odisha',
    'assam', 'jharkhand', 'uttarakhand', 'goa', 'bihar', 'chhattisgarh',
    # Other common countries / regions
    'usa', 'uk', 'canada', 'australia', 'singapore', 'germany',
    'california', 'texas', 'york', 'london', 'remote',
}


def _is_name_word(w: str) -> bool:
    """Return True if a word looks like part of a person's name."""
    w_clean = w.rstrip('.')
    if len(w_clean) == 1:
        return w_clean.isupper()  # single uppercase initial (e.g. "K" in "Mounish K")
    # ALL CAPS surname or initial block like "TEJASWI", "KUMAR", "N"
    if w_clean.isupper() and w_clean.isalpha():
        return w_clean.lower() not in _NON_NAME_WORDS
    # Standard Title Case word like "Smith", "Nair", "Mary-Jane"
    return bool(re.match(r'^[A-Z][a-z\-]+$', w_clean))


def _check_line_for_name(line: str) -> Optional[str]:
    """Return a name if `line` looks like a candidate name, else None."""
    if not line or any(c.isdigit() for c in line):
        return None
    if any(tok in line.lower() for tok in ('@', '\\', 'http', 'linkedin', 'github', 'www.')):
        return None

    # Handle "Label: value" patterns — e.g. "Name: Rahul Kumar", "Candidate Name: Priya Sharma"
    # Strip the label prefix and try to parse the value portion as a name.
    label_match = re.match(r'^[A-Za-z\s]{1,20}:\s*(.+)$', line.strip())
    if label_match:
        value_part = label_match.group(1).strip()
        # Only recurse if value_part has no colon (avoid double-recursion)
        if ':' not in value_part and value_part:
            return _check_line_for_name(value_part)
        return None

    # Strip non-name characters to get clean words first
    clean_line = re.sub(r'[^a-zA-Z\s.\-]', '', line).strip().strip('.')
    words = clean_line.split()
    if not (2 <= len(words) <= 6):
        return None
    if clean_line.lower().rstrip('.') in _SECTION_HEADER_BLACKLIST:
        return None

    # Run blacklist check BEFORE comma-reversed pattern — this prevents
    # "Hyderabad, Telangana" or "City, State" address lines from being
    # incorrectly treated as a reversed name ("Telangana Hyderabad").
    lower_words = {w.lower().rstrip('.') for w in words}
    if lower_words & _NON_NAME_WORDS:
        return None

    # Normalise comma-reversed names: "Smith, John" → "John Smith"
    # Only after confirming neither word is a geographic/non-name term above.
    comma_match = re.match(r'^([A-Z][A-Za-z\-]+),\s+([A-Z][A-Za-z\-]+(?:\s+[A-Z][A-Za-z\-]+)?)$', line.strip())
    if comma_match:
        candidate = f"{comma_match.group(2)} {comma_match.group(1)}"
        if candidate.lower() not in _SECTION_HEADER_BLACKLIST:
            return candidate

    # Title Case + initials check — handles:
    #   "John Smith", "J. Smith", "Mounish K", "N. V. S. Tejaswi", "N. V. S. TEJASWI"
    if all(_is_name_word(w) for w in words):
        non_initial = [w for w in words if len(w.rstrip('.')) > 1]
        if len(non_initial) >= 1:
            # Normalise: Title Case everything; preserve dots on initials
            return ' '.join(w[0].upper() + w[1:].lower() if (len(w.rstrip('.')) > 1 and w.rstrip('.').isupper()) else w for w in words)

    # ALL CAPS check — e.g. "RAHUL NAIR", "N V S TEJASWI" (after dot-stripping)
    alpha_only = re.sub(r'[^a-zA-Z\s]', '', clean_line).strip()
    alpha_words = alpha_only.split()
    if 2 <= len(alpha_words) <= 5 and alpha_only.isupper():
        alpha_lower = {w.lower() for w in alpha_words}
        if not (alpha_lower & _NON_NAME_WORDS) and alpha_only.lower() not in _SECTION_HEADER_BLACKLIST:
            return alpha_only.title()

    return None


def is_plausible_name(name: str) -> bool:
    """Return False if name contains geographic or non-name blacklisted words.

    Used to validate LLM-extracted names before trusting them.
    """
    if not name or not name.strip():
        return False
    clean = re.sub(r'[^a-zA-Z\s]', '', name).strip().lower()
    words = clean.split()
    if not words:
        return False
    if set(words) & _NON_NAME_WORDS:
        return False
    # Must have at least one word that looks like a real name word (title-cased check)
    return any(_is_name_word(w.title()) for w in words)


def extract_name_from_pdf_fonts(path: str) -> Optional[str]:
    """Extract the candidate name using font-size analysis (primary ATS method).

    Real-world ATS (Greenhouse, Lever, Workday) rely on font-size signals: the
    candidate's name is almost always the largest text in the top portion of the
    first page.  This uses PyMuPDF span-level dicts to read font sizes per text
    fragment, so two-column PDFs or address-before-name layouts don't confuse it.

    Returns None for non-PDF paths, if fitz is unavailable, or if no plausible
    name is found among the largest-font spans.
    """
    if not path.lower().endswith(".pdf"):
        return None
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return None

    try:
        doc = fitz.open(path)
        if not doc.page_count:
            doc.close()
            return None

        page = doc[0]
        page_height = page.rect.height
        top_zone_limit = page_height * 0.30  # examine only the top 30% of page 1

        span_data: list[dict] = []
        for block in page.get_text("dict").get("blocks", []):
            if block.get("type") != 0:  # skip image blocks
                continue
            for line in block.get("lines", []):
                line_y0 = line["bbox"][1]
                if line_y0 > top_zone_limit:
                    continue
                for span in line.get("spans", []):
                    text = span.get("text", "").strip()
                    if not text:
                        continue
                    span_data.append({
                        "size": float(span.get("size", 0)),
                        "y":    float(line["bbox"][1]),
                        "x":    float(span["bbox"][0]),
                        "text": text,
                    })

        doc.close()

        if not span_data:
            return None

        max_size = max(s["size"] for s in span_data)
        if max_size < 8:  # ignore pathologically small text
            return None

        # Collect spans within 1.5pt of the maximum font size, sorted top-left→right
        large_spans = sorted(
            [s for s in span_data if s["size"] >= max_size - 1.5],
            key=lambda s: (round(s["y"] / 4) * 4, s["x"]),
        )

        # Group into visual lines by y-proximity (within 4pt ≈ same line)
        lines: list[str] = []
        current_bucket: Optional[int] = None
        current_parts: list[str] = []
        for s in large_spans:
            bucket = round(s["y"] / 4)
            if current_bucket is None or bucket != current_bucket:
                if current_parts:
                    lines.append(" ".join(current_parts))
                current_bucket = bucket
                current_parts = [s["text"]]
            else:
                current_parts.append(s["text"])
        if current_parts:
            lines.append(" ".join(current_parts))

        # Try each reconstructed line through the name validator
        for line_text in lines:
            result = _check_line_for_name(line_text)
            if result:
                return result

        # Last attempt: join all large-font lines (handles names split across spans)
        if len(lines) > 1:
            result = _check_line_for_name(" ".join(lines))
            if result:
                return result

        return None

    except Exception as exc:
        logger.debug("Font-size name extraction failed for %s: %s", path, exc)
        return None


def extract_name(text: str) -> Optional[str]:
    """Return the candidate's name by examining only the resume header.

    Names appear on the FIRST line (or first two lines) of every resume.
    We look at the first 6 non-empty lines only — scanning further risks
    matching city/state address lines, job titles, or institution names.

    Handles: Title Case, ALL CAPS, mixed-initials ("N. V. S. TEJASWI"),
    comma-reversed ("Smith, John"), and pipe-separated contact headers
    ("Siddhi Borawake | +91 … | email@gmail.com").
    """
    non_empty_lines = [l.strip() for l in text.splitlines() if l.strip()]

    for raw_line in non_empty_lines[:6]:
        # Pipe-separated header: try only the first segment (the name part)
        if '|' in raw_line:
            first_segment = raw_line.split('|')[0].strip()
            result = _check_line_for_name(first_segment)
            if result:
                return result
            # Entire line is a contact-info row — skip it, don't fall through
            continue

        result = _check_line_for_name(raw_line)
        if result:
            return result

    return None


class NeedsOCRError(ValueError):
    """Raised when a PDF yields too little text — likely a scanned image."""


@dataclass
class ParsedDocument:
    text: str
    page_count: int
    source_format: str                        # "pdf" | "docx"
    metadata: Dict[str, Any] = field(default_factory=dict)


def _find_column_split(blocks: list, page_width: float) -> Optional[float]:
    """Return the x-coordinate of a two-column split, or None for single-column pages.

    Uses a bucket histogram of block left-edges; a near-zero bucket in the
    30–70% page-width zone indicates a gutter between columns.
    """
    if page_width <= 0 or len(blocks) < 6:
        return None
    x_starts = [b[0] for b in blocks]
    bucket_size = max(page_width / 20, 1.0)
    buckets = [0] * 20
    for x in x_starts:
        buckets[min(int(x / bucket_size), 19)] += 1
    lo_b = int(page_width * 0.30 / bucket_size)
    hi_b = min(int(page_width * 0.70 / bucket_size), 19)
    center = buckets[lo_b: hi_b + 1]
    if not center or min(center) > 1:
        return None  # no clear gutter
    gap_bucket = lo_b + center.index(min(center))
    split = (gap_bucket + 0.5) * bucket_size
    # Both sides must actually have blocks — otherwise it's a wide single column
    if not any(x < split for x in x_starts) or not any(x >= split for x in x_starts):
        return None
    return split


def _blocks_in_reading_order(blocks: list, page_width: float) -> list:
    """Sort text blocks in natural reading order.

    Single-column: top-to-bottom (y-bin / x).
    Two-column: left column top-to-bottom, then right column top-to-bottom.
    """
    text_blocks = [b for b in blocks if b[4].strip()]
    if not text_blocks:
        return text_blocks
    split = _find_column_split(text_blocks, page_width)
    if split is None:
        return sorted(text_blocks, key=lambda b: (round(b[1] / 15) * 15, b[0]))
    left = sorted([b for b in text_blocks if b[0] < split], key=lambda b: b[1])
    right = sorted([b for b in text_blocks if b[0] >= split], key=lambda b: b[1])
    return left + right


def parse_pdf(path: str, ocr_fallback: bool = True) -> ParsedDocument:
    """Extract text from a PDF using PyMuPDF.

    Handles single-column and two-column resume layouts by detecting a gutter
    between columns and sorting each column independently before joining.

    Raises:
        NeedsOCRError: When total extracted text is fewer than 100 characters.
    """
    import fitz  # PyMuPDF

    doc = fitz.open(path)
    page_texts: List[str] = []
    uri_links: List[str] = []

    for page in doc:
        # blocks → (x0, y0, x1, y1, text, block_no, block_type)
        blocks = page.get_text("blocks")
        blocks_sorted = _blocks_in_reading_order(blocks, page.rect.width)
        page_text = "\n".join(b[4].strip() for b in blocks_sorted if b[4].strip())
        page_texts.append(page_text)

        # Collect hyperlink annotations (URLs hidden behind display text like "GitHub")
        for link in page.get_links():
            uri = link.get("uri", "")
            if uri and uri.startswith("http"):
                uri_links.append(uri)

    full_text = "\n".join(page_texts).strip()
    # Append extracted URIs so regex-based link extraction can find them
    if uri_links:
        full_text += "\n" + " ".join(uri_links)
    doc.close()

    if len(full_text) < 100:
        if ocr_fallback:
            ocr_text = _ocr_pdf(path)
            if ocr_text and len(ocr_text) >= 100:
                return ParsedDocument(
                    text=ocr_text,
                    page_count=len(page_texts),
                    source_format="pdf",
                    metadata={"path": path, "ocr": True},
                )
        raise NeedsOCRError(
            "Extracted text is too short (< 100 chars). "
            "This PDF may be scanned and OCR could not extract sufficient text."
        )

    return ParsedDocument(
        text=full_text,
        page_count=len(page_texts),
        source_format="pdf",
        metadata={"path": path},
    )


def _ocr_pdf(path: str) -> str:
    """Render each PDF page as an image and extract text via Tesseract OCR.

    Returns extracted text, or empty string if pytesseract/tesseract is unavailable.
    """
    try:
        import fitz
        import pytesseract
        from PIL import Image
        import io
    except ImportError:
        return ""

    try:
        doc = fitz.open(path)
        texts: List[str] = []
        for page in doc:
            # Render at 2× DPI for better OCR accuracy
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            page_text = pytesseract.image_to_string(img, lang="eng")
            if page_text.strip():
                texts.append(page_text.strip())
        doc.close()
        return "\n".join(texts)
    except Exception:
        return ""


def parse_docx(path: str) -> ParsedDocument:
    """Extract text from a DOCX file using python-docx.

    Iterates paragraphs (preserving style names as metadata hints) and all table cells.
    """
    from docx import Document  # python-docx

    doc = Document(path)
    parts: List[str] = []
    style_hints: List[str] = []

    # Body paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
            if para.style and para.style.name:
                style_hints.append(para.style.name)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    parts.append(cell_text)

    full_text = "\n".join(parts).strip()

    return ParsedDocument(
        text=full_text,
        page_count=1,   # DOCX has no reliable page count; default to 1
        source_format="docx",
        metadata={"path": path, "style_hints": style_hints},
    )


def parse_resume(path: str, ocr_fallback: bool = True) -> ParsedDocument:
    """Parse a PDF or DOCX resume file and return a ParsedDocument.

    Dispatches to ``parse_pdf`` or ``parse_docx`` based on file extension.

    Args:
        path: Absolute path to the uploaded resume file.
        ocr_fallback: When True, attempt OCR for scanned PDFs. Default True.

    Returns:
        ParsedDocument with extracted text, page count, format, and metadata.

    Raises:
        ValueError: If the file extension is not .pdf or .docx/.doc.
        NeedsOCRError: If the PDF appears to be a scanned image.
    """
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return parse_pdf(path, ocr_fallback=ocr_fallback)
    elif ext in {".docx", ".doc"}:
        return parse_docx(path)
    else:
        raise ValueError(
            f"Unsupported file format: {ext!r}. Only PDF and DOCX are accepted."
        )


# Backwards-compatible alias used in existing router code
parse_document = parse_resume


def extract_metadata_via_llm(text: str) -> Optional[dict]:
    """Use the configured LLM (Groq/OpenAI) to extract structured metadata from resume text.

    Returns a dict with keys: name, email, phone, title, experience_level,
    years_experience, graduation_year, linkedin_url, github_url, portfolio_url.
    All values may be None.  Returns None if LLM is not configured or the call fails.
    """
    try:
        from app.config import settings
        if not settings.llm_api_key:
            return None
    except Exception:
        return None

    try:
        from openai import OpenAI
        client = OpenAI(
            api_key=settings.llm_api_key,
            base_url=settings.llm_base_url,
            timeout=float(settings.llm_timeout),
        )

        # Send only the first 300 chars as a dedicated header block so the
        # LLM focuses on the name at the top, then the full text for context.
        header_block = text[:300].strip()
        sample = text[:8000]
        prompt = (
            "Extract structured information from this resume.\n"
            "IMPORTANT — for the 'name' field: the candidate's full name is ALWAYS "
            "the very first prominent line at the top of the resume, before any "
            "contact details, taglines, or section headings. It may be in ALL CAPS "
            "or include initials with dots (e.g. 'N. V. S. Tejaswi', 'RAHUL KUMAR'). "
            "Do NOT return a city, state, job title, university, or tagline as the name.\n\n"
            f"Resume header (first 300 chars — name is here):\n{header_block}\n\n"
            "Return ONLY valid JSON with these exact keys (use null for missing fields):\n"
            '{"name":null,"email":null,"phone":null,"title":null,'
            '"experience_level":null,"years_experience":null,"graduation_year":null,'
            '"linkedin_url":null,"github_url":null,"portfolio_url":null}\n\n'
            "experience_level must be one of: entry, junior, mid, senior, executive\n"
            "entry=0-1 yr (interns, fresh grads, trainees). "
            "junior=1-3 yrs. mid=3-5 yrs. senior=5-10 yrs. executive=10+ yrs or C-level/VP/Director.\n"
            "years_experience counts ONLY full-time/permanent roles. "
            "Internship, trainee, and student positions do NOT count toward years_experience.\n"
            "If the candidate only has internship experience, set years_experience=null.\n"
            "graduation_year is a 4-digit integer.\n\n"
            f"Full resume:\n{sample}"
        )

        response = client.chat.completions.create(
            model=settings.llm_model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
            max_tokens=400,
        )

        content = (response.choices[0].message.content or "").strip()
        # Strip markdown code fences if present
        content = re.sub(r"^```[a-z]*\n?", "", content)
        content = re.sub(r"\n?```$", "", content).strip()

        json_match = re.search(r"\{.*\}", content, re.DOTALL)
        if not json_match:
            return None

        data: dict = json.loads(json_match.group())

        # Sanitise experience_level
        valid_levels = {"entry", "junior", "mid", "senior", "executive"}
        if data.get("experience_level") not in valid_levels:
            data["experience_level"] = None

        # Sanitise years_experience
        yoe = data.get("years_experience")
        if yoe is not None:
            try:
                yoe = float(yoe)
                data["years_experience"] = yoe if 0 < yoe < 50 else None
            except (TypeError, ValueError):
                data["years_experience"] = None

        # Sanitise graduation_year
        gy = data.get("graduation_year")
        if gy is not None:
            try:
                gy = int(gy)
                data["graduation_year"] = gy if 1970 <= gy <= datetime.now().year + 2 else None
            except (TypeError, ValueError):
                data["graduation_year"] = None

        logger.debug("LLM metadata extracted: name=%s email=%s", data.get("name"), data.get("email"))
        return data

    except Exception as exc:
        logger.debug("LLM metadata extraction failed: %s", exc)
        return None
